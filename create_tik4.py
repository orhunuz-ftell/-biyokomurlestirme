# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FEN BILIMLERI ENSTITUSU\n')
run.bold = True
run.font.size = Pt(14)
run = title.add_run('DOKTORA TEZ IZLEME KOMITESI TOPLANTISI-4\n')
run.bold = True
run.font.size = Pt(14)
run = title.add_run('ILERLEME RAPORU')
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

# Student info
info = doc.add_paragraph()
info.add_run('Ogrenci Adi Soyadi: ').bold = True
info.add_run('Orhun UZDIYEM\n')
info.add_run('Ogrenci No: ').bold = True
info.add_run('[Ogrenci Numaraniz]\n')
info.add_run('Anabilim Dali: ').bold = True
info.add_run('Kimya Muhendisligi\n')
info.add_run('Programi: ').bold = True
info.add_run('Doktora\n')
info.add_run('Danisman: ').bold = True
info.add_run('[Danisman Adi]\n')
info.add_run('Tez Konusu: ').bold = True
info.add_run('Biyokutle Pirolizi Biyoyag Uretimi ve Ters Makine Ogrenmesi ile Tahmin\n')
info.add_run('Rapor Tarihi: ').bold = True
info.add_run(f'{datetime.now().strftime("%d.%m.%Y")}\n')

doc.add_paragraph()

# Section 1
heading = doc.add_heading('1. ONCEKI DONEM CALISMALARI (TIK-3\'ten Sonra)', level=1)

para = doc.add_paragraph(
    'TIK-3 toplantisind an sonra asagidaki calismalar tamamlanmistir:'
)

# Subsection 1.1
doc.add_heading('1.1. Kimyasal Proses Secimi ve Modelleme Yaklasimi', level=2)

para = doc.add_paragraph(
    'Literatur taramasi sonucunda biyoyagdan hidrojen uretimi icin buhar reforming prosesi secilmistir. '
    'Bu proses 5 alternatif arasinda degerlendirilerek en yuksek skoru (%93) almistir. '
    'Secim kriterleri: Modelleme karmasikligi, veri mevcudiyeti, simulasyon hizi, endustriyel onemi, '
    'makine ogrenmesi uygunlugu, olceklenebilirlik, akademik deger ve uygulanabilirlik.'
)

para = doc.add_paragraph()
para.add_run('Aspen Plus lisans sorunu nedeniyle acik kaynak ')
run = para.add_run('Cantera')
run.bold = True
para.add_run(' termodinamik simulasyon yazilimi kullanilmistir.')

doc.add_heading('1.2. Ozel Cantera Mekanizmasi Gelistirilmesi', level=2)

para = doc.add_paragraph(
    'GRI-Mech 3.0 mekanizmasi (53 tur, 325 reaksiyon) biyoyag icin uygun olmadigindan, '
    'ozel bir termodinamik mekanizma gelistirilmistir:'
)

para = doc.add_paragraph()
para.add_run('Biyoyag Surrogate Turleri:\n').bold = True

surrogate_table = doc.add_table(rows=7, cols=3)
surrogate_table.style = 'Light Grid Accent 1'

# Header
header_cells = surrogate_table.rows[0].cells
header_cells[0].text = 'Biyoyag Grubu'
header_cells[1].text = 'Surrogate Tur'
header_cells[2].text = 'Kimyasal Formul'

# Data
data = [
    ['Aromatikler', 'Toluen', 'C7H8'],
    ['Asitler', 'Asetik asit', 'CH3COOH'],
    ['Alkoller', 'Etanol', 'C2H5OH'],
    ['Furanlar', 'Furan', 'C4H4O'],
    ['Fenoller', 'Fenol', 'C6H6O'],
    ['Aldehit/Keton', 'Aseton', 'C3H6O']
]

for i, row_data in enumerate(data, 1):
    row_cells = surrogate_table.rows[i].cells
    for j, cell_data in enumerate(row_data):
        row_cells[j].text = cell_data

doc.add_paragraph()
para = doc.add_paragraph()
para.add_run('Toplam: ').bold = True
para.add_run('59 kimyasal tur (6 biyoyag + 53 GRI-Mech)')

doc.add_heading('1.3. Veritabani Semasi Guncellemesi', level=2)

para = doc.add_paragraph(
    'SQL Server veritabani (BIOOIL) Cantera simulasyonlari icin genisletilmistir:'
)

items = [
    'AspenSimulation tablosuna Temperature_C, Pressure_bar, SC_ratio sutunlari eklendi',
    'SimulationSource sutunu eklendi ("Cantera" tanimlayicisi icin)',
    'ConvergenceStatus sutunu guncellendi',
    'ReformingConditions, HydrogenProduct, SyngasComposition, EnergyBalance tablolari mevcut'
]

for item in items:
    doc.add_paragraph(item, style='List Bullet')

# Section 2
doc.add_heading('2. TAMAMLANAN SIMULASYONLAR VE VERI URETIMI', level=1)

doc.add_heading('2.1. Simulasyon Matrisi', level=2)

para = doc.add_paragraph(
    'Toplam 1.170 hidrojen uretim senaryosu olusturulmustur:'
)

para = doc.add_paragraph()
para.add_run('26 biyoyag kompozisyonu').bold = True
para.add_run(' x ')
para.add_run('45 proses kosulu').bold = True
para.add_run(' = ')
para.add_run('1.170 senaryo').bold = True

doc.add_paragraph()

# Process conditions table
proc_table = doc.add_table(rows=4, cols=3)
proc_table.style = 'Light Grid Accent 1'

header = proc_table.rows[0].cells
header[0].text = 'Parametre'
header[1].text = 'Aralik'
header[2].text = 'Seviye Sayisi'

proc_data = [
    ['Sicaklik', '650-850 C', '5 (650, 700, 750, 800, 850)'],
    ['Basinc', '5-30 bar', '3 (5, 17.5, 30)'],
    ['S/C Orani', '2-6', '3 (2, 4, 6)']
]

for i, row_data in enumerate(proc_data, 1):
    cells = proc_table.rows[i].cells
    for j, val in enumerate(row_data):
        cells[j].text = val

doc.add_heading('2.2. Modellenen Proses Asamalari', level=2)

stages = [
    ('Buhar Reforming', 'Biyoyag + H2O -> Sentez gazi (H2, CO, CO2, CH4)', 'Gibbs serbest enerji minimizasyonu'),
    ('Yuksek Sicaklik Shift (370 C)', 'CO + H2O -> CO2 + H2', 'H2 artisi, CO azalmasi'),
    ('Dusuk Sicaklik Shift (210 C)', 'CO + H2O -> CO2 + H2', 'Ilave H2 zenginlestirme'),
    ('Flash Ayirma (40 C)', 'Su giderimi', '%99.9 su giderme verimliligi'),
    ('CO2 Giderimi', 'Kimyasal absorpsiyon', '%95 CO2 giderme verimliligi'),
    ('PSA (25 bar)', 'H2 saflastirma', '%99.9 H2 safligi, %88 H2 geri kazanimi')
]

for stage, desc, note in stages:
    para = doc.add_paragraph()
    para.add_run(f'{stage}: ').bold = True
    para.add_run(f'{desc} ({note})')

doc.add_heading('2.3. Simulasyon Sonuclari', level=2)

para = doc.add_paragraph()
para.add_run('Tum simulasyonlar basariyla tamamlanmistir:\n').bold = True

# Results table
results_table = doc.add_table(rows=6, cols=2)
results_table.style = 'Light Grid Accent 1'

results_data = [
    ['Toplam Simulasyon', '1.170'],
    ['Basarili', '1.170 (%100)'],
    ['Basarisiz', '0'],
    ['Calisma Suresi', '8.7 saniye'],
    ['Ortalama Sure/Simulasyon', '0.01 saniye'],
    ['Veritabanina Kaydedilen', '1.170 kayit']
]

for i, (key, val) in enumerate(results_data):
    cells = results_table.rows[i].cells
    cells[0].text = key
    cells[1].text = val

doc.add_paragraph()

para = doc.add_paragraph()
para.add_run('Veri Dagilimi:\n').bold = True

dist_items = [
    'Sicaklik: Her seviyede 234 simulasyon (650, 700, 750, 800, 850 C)',
    'Basinc: Her seviyede 390 simulasyon (5, 17.5, 30 bar)',
    'S/C Orani: Her seviyede 390 simulasyon (2, 4, 6)'
]

for item in dist_items:
    doc.add_paragraph(item, style='List Bullet')

# Section 3
doc.add_heading('3. SISTEM MIMARISI VE YAZILIM GELISTIRME', level=1)

para = doc.add_paragraph(
    'Python tabanli moduler bir sistem gelistirilmistir (7 ana modul):'
)

modules = [
    ('cantera_input_processor.py', '1.170 senaryoyu yukler, Cantera girislerini hazirlar'),
    ('cantera_equilibrium.py', 'Gibbs minimizasyonu (reformer, HTS, LTS)'),
    ('separation_models.py', 'Ayirma prosesleri (Flash, CO2 giderimi, PSA)'),
    ('property_calculator.py', '16 makine ogrenmesi ozelligi hesaplar'),
    ('database_writer.py', 'SQL Server veritabanina yazar'),
    ('validation.py', '5 seviyeli dogrulama sistemi'),
    ('generate_data_cantera.py', 'Ana kontrol scripti')
]

for module, desc in modules:
    para = doc.add_paragraph()
    para.add_run(f'{module}: ').bold = True
    para.add_run(desc)

doc.add_heading('3.1. Dogrulama Sistemi', level=2)

para = doc.add_paragraph('5 seviyeli dogrulama uygulanmistir:')

validation_levels = [
    'Seviye 1: Kutle ve enerji dengesi (mol kesirleri toplami = 1.0)',
    'Seviye 2: Fiziksel araliklar (H2 verimi 5-15 kg/100kg biyoyag)',
    'Seviye 3: Termodinamik uygunluk (H2 artisi WGS reaktorlerinde)',
    'Seviye 4: Istatistiksel tutarlilik (literatur ile karsilastirma)',
    'Seviye 5: Makine ogrenmesi hazirligi (tum 16 ozellik mevcut, NaN yok)'
]

for level in validation_levels:
    doc.add_paragraph(level, style='List Bullet')

# Section 4
doc.add_heading('4. ELDE EDILEN VERI SETI OZELLIKLERI', level=1)

para = doc.add_paragraph()
para.add_run('Makine ogrenmesi icin hazir 1.170 kayitlik veri seti:\n').bold = True

para = doc.add_paragraph()
para.add_run('Girdiler (Biyoyag Kompozisyonu):\n').bold = True

inputs = ['Aromatikler (%)', 'Asitler (%)', 'Alkoller (%)', 'Furanlar (%)', 'Fenoller (%)', 'Aldehit/Keton (%)']
for inp in inputs:
    doc.add_paragraph(inp, style='List Bullet 2')

para = doc.add_paragraph()
para.add_run('Proses Kosullari:\n').bold = True
conditions = ['Sicaklik (C)', 'Basinc (bar)', 'S/C orani']
for cond in conditions:
    doc.add_paragraph(cond, style='List Bullet 2')

para = doc.add_paragraph()
para.add_run('Ciktilar (16 Ozellik):\n').bold = True
outputs = [
    'H2 verimi (kg/100kg biyoyag)',
    'H2 safligi (%)',
    'H2 uretim hizi (mol/s)',
    'Karbon donusumu (%)',
    'Enerji verimliligi (%)',
    'H2/CO orani',
    'Sentez gazi kompozisyonu (H2, CO, CO2, CH4)',
    'Urun kompozisyonu',
    'Su tuketimi',
    'Gercek S/C orani',
    'Toplam H2 geri kazanimi'
]
for out in outputs:
    doc.add_paragraph(out, style='List Bullet 2')

# Section 5
doc.add_heading('5. BILIMSEL KATKI VE YENILIK', level=1)

contributions = [
    'Biyoyag buhar reforming icin ozel Cantera mekanizmasi gelistirilmesi (literaturde ilk)',
    'Acik kaynak yazilim kullanarak ticari simulator seviyesinde veri uretimi',
    'Hizli veri uretimi: 1.170 simulasyon 8.7 saniyede (Aspen Plus ile saatler surerdi)',
    'Moduler ve yeniden kullanilabilir Python kutuphanesi',
    'Tam otomatik veri uretim pipeline',
    'Kapsamli 5 seviyeli dogrulama sistemi'
]

for contrib in contributions:
    doc.add_paragraph(contrib, style='List Bullet')

# Section 6
doc.add_heading('6. SINIRLAMALAR VE GECERLILIK', level=1)

para = doc.add_paragraph('Gelistirilen sistemin sinirlam alari:')

limitations = [
    'Termodinamik denge varsayimi (kinetik modellenmemistir)',
    'Basitlestirilmis PSA modeli (detayli adsorpsiyon dinamigi yok)',
    'Biyoyag surrogate turleri icin yaklasik termodinamik data',
    'Ticari simulatorlere gore beklenen dogruluk: %75-85'
]

for lim in limitations:
    doc.add_paragraph(lim, style='List Bullet')

para = doc.add_paragraph()
para.add_run('\nBu sinirlamalar ')
para.add_run('makine ogrenmesi egitim verisi').bold = True
para.add_run(' uretimi icin kabul edilebilir seviyededir ve tez kapsaminda belgelenecektir.')

# Section 7
doc.add_heading('7. SONRAKI ADIMLAR (Faz 4: Makine Ogrenmesi)', level=1)

para = doc.add_paragraph('Veri seti hazir oldugu icin bir sonraki asamaya gecilecektir:')

next_steps = [
    'Kesifsel veri analizi (EDA) ve veri gorsellestirme',
    'Ters makine ogrenmesi modeli gelistirme (H2 ozellikleri -> Biyoyag kompozisyonu)',
    'Model secimi ve karsilastirma (Random Forest, Neural Network, XGBoost)',
    'Hiperparametre optimizasyonu',
    'Capraz dogrulama ve performans metrikleri',
    'Duyarlilik analizi',
    'Literatur ile karsilastirma ve validasyon',
    'Tez yazimi'
]

for step in next_steps:
    doc.add_paragraph(step, style='List Number')

# Section 8
doc.add_heading('8. ZAMAN CIZELGESI', level=1)

timeline_table = doc.add_table(rows=7, cols=3)
timeline_table.style = 'Light Grid Accent 1'

header = timeline_table.rows[0].cells
header[0].text = 'Ay'
header[1].text = 'Faaliyet'
header[2].text = 'Durum'

timeline_data = [
    ['Ocak-Subat 2025', 'Veri analizi ve on isleme', 'Planlaniy or'],
    ['Mart 2025', 'ML model gelistirme ve egitim', 'Planlaniyor'],
    ['Nisan 2025', 'Model optimizasyonu ve validasyon', 'Planlaniyor'],
    ['Mayis 2025', 'Sonuclarin yorumlanmasi ve analiz', 'Planlaniyor'],
    ['Haziran 2025', 'Tez yazimi - Bolum 1-3', 'Planlaniyor'],
    ['Temmuz-Agustos 2025', 'Tez yazimi - Bolum 4-6 ve tamamlama', 'Planlaniyor']
]

for i, (month, activity, status) in enumerate(timeline_data, 1):
    cells = timeline_table.rows[i].cells
    cells[0].text = month
    cells[1].text = activity
    cells[2].text = status

# Section 9
doc.add_heading('9. YAYINLAR VE SUNUMLAR', level=1)

para = doc.add_paragraph()
para.add_run('Planlanan Yayinlar:\n').bold = True

publications = [
    'Cantera mekanizmasi gelistirme ve biyoyag modellemesi (journal makale)',
    'Ters makine ogrenmesi yaklasimi (journal makale)',
    'Konferans sunumlari (ulusal/uluslararasi)'
]

for pub in publications:
    doc.add_paragraph(pub, style='List Bullet')

# Section 10
doc.add_heading('10. SONUC', level=1)

conclusion_paras = [
    'Bu donemde buyuk bir ilerleme kaydedilmistir. Aspen Plus lisans sorunu basariyla '
    'acik kaynak Cantera yazilimi ile asilmis ve orijinal bir cozum gelistirilmistir.',

    'Toplam 1.170 hidrojen uretim simulasyonu basariyla tamamlanmis ve SQL Server '
    'veritabanina kaydedilmistir. Veri seti makine ogrenmesi icin tamamen hazirdir.',

    'Gelistirilen sistem moduler, yeniden kullanilabilir ve bilimsel acidan gecerlidir. '
    'Tum kod ve dokumantasyon GitHub deposunda saklanmaktadir.',

    'Bir sonraki TIK toplantisina kadar makine ogrenmesi modelleri gelistirilecek ve '
    'sonuclar degerlendirilecektir.'
]

for para_text in conclusion_paras:
    doc.add_paragraph(para_text)

doc.add_paragraph()
doc.add_paragraph()

# Signature section
signature = doc.add_paragraph()
signature.add_run('\nOgrenci Adi Soyadi ve Imzasi\n\n\n')
signature.add_run('_______________________________\n')
signature.add_run('Orhun UZDIYEM')

signature = doc.add_paragraph()
signature.add_run('\nDanisman Adi Soyadi ve Imzasi\n\n\n')
signature.add_run('_______________________________\n')
signature.add_run('[Danisman Adi]')

# Save document
doc.save('OrhunUzdiyem_tik4.docx')
print('[OK] TIK-4 report created successfully: OrhunUzdiyem_tik4.docx')
